"""
6_CIP_Form.py — NYS OCM Community Impact Plan (CIP) Form Builder
Guided form that generates a submission-ready .docx file.
"""

import sys, os, io
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import streamlit as st
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

st.set_page_config(
    page_title="CIP Form | NYS Cannabis Tool",
    page_icon="📋",
    layout="wide",
)

st.markdown("""
<style>
.disclaimer-box {
    background: #fff3cd; border: 2px solid #e0a800;
    border-radius: 8px; padding: 16px 20px; margin-bottom: 20px; font-size: 0.92rem;
}
.section-note {
    background: #e8f4fd; border-left: 4px solid #1565C0;
    padding: 10px 14px; border-radius: 4px;
    font-size: 0.88rem; color: #0d3349; margin-bottom: 12px;
}
.char-warn { color: #856404; font-size: 0.82rem; }
.char-over { color: #721c24; font-size: 0.82rem; font-weight: bold; }
</style>
""", unsafe_allow_html=True)

# ── Document builder ───────────────────────────────────────────────────────────

def _add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1b, 0x5e, 0x20)


def _add_question(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def _add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)


def _add_field_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{label}  ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    val_run = p.add_run(value or "")
    val_run.font.size = Pt(11)


def _v(d, key, fallback="[not provided]"):
    val = d.get(key, "")
    return val.strip() if val and val.strip() else fallback


def build_cip_docx(d):
    """Build and return a CIP .docx as bytes from the form data dict d."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1 = title.add_run("NEW YORK STATE OFFICE OF CANNABIS MANAGEMENT\n")
    t1.bold = True
    t1.font.size = Pt(13)
    t2 = title.add_run("Community Impact Plan")
    t2.bold = True
    t2.font.size = Pt(13)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = subtitle.add_run(_v(d, "business_name") + "  ·  Submitted " + _v(d, "cert_date"))
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # ── SECTION 1 ──────────────────────────────────────────────────────────────
    _add_section_title(doc, "SECTION 1: Applicant Information")
    for label, key in [
        ("Business / Applicant Name:", "business_name"),
        ("License Type Applied For:", "license_type"),
        ("Primary Contact Name:", "contact_name"),
        ("Title / Role:", "title_role"),
        ("Address:", "address"),
        ("City, State, ZIP:", "city_state_zip"),
        ("County:", "county"),
        ("OCM Region:", "ocm_region"),
        ("Phone:", "phone"),
        ("Email:", "email"),
    ]:
        _add_field_line(doc, label, _v(d, key, ""))

    # ── SECTION 2 ──────────────────────────────────────────────────────────────
    _add_section_title(doc, "SECTION 2: Problem / Need Statement")

    _add_question(doc, "2a. Describe the problem or need your Community Impact Plan is designed to address.")
    _add_body(doc,
        f"{_v(d,'business_name')} is a {_v(d,'business_type')} located in {_v(d,'county')} County, "
        f"New York, within the {_v(d,'ocm_region')} OCM region. Like communities across New York State, "
        f"{_v(d,'county')} County has a history of cannabis prohibition that disproportionately affected "
        f"individuals from {_v(d,'affected_backgrounds')} backgrounds, particularly those in low-income or "
        f"under-resourced neighborhoods. Despite the legalization of adult-use cannabis, many of the "
        f"communities most harmed by prior enforcement lack access to {_v(d,'lack_access_to')}. Our business "
        f"recognizes that meaningful legalization must include intentional investment in these communities. "
        f"Through this Community Impact Plan, we aim to address {_v(d,'unmet_need')} in "
        f"{_v(d,'region_name')} and the surrounding region."
    )

    _add_question(doc, "2b. Who are the individuals or groups most affected by this problem?")
    _add_body(doc,
        f"The individuals most affected by the legacy of cannabis prohibition in our area include "
        f"{_v(d,'most_affected_description')}. Based on our conversations with {_v(d,'partners_consulted')} "
        f"and our own knowledge of the region, we estimate that {_v(d,'estimated_people')} people in "
        f"{_v(d,'benefit_area')} could benefit from the programming described in this plan. We intend to "
        f"prioritize outreach to {_v(d,'priority_populations')} and will partner with "
        f"{_v(d,'outreach_partner')} to help us reach individuals who may not otherwise engage with the "
        f"legal cannabis industry."
    )

    # ── SECTION 3 ──────────────────────────────────────────────────────────────
    _add_section_title(doc, "SECTION 3: Research, Evidence Base, and Partnerships")

    _add_question(doc, "3a. What research or evidence informed your Community Impact Plan?")
    _add_body(doc,
        f"In developing this Community Impact Plan, {_v(d,'business_name')} drew on several sources of "
        f"information. We reviewed {_v(d,'sources_reviewed')} and consulted with {_v(d,'orgs_consulted')}. "
        f"Cornell Cooperative Extension provided access to research and best practices related to "
        f"{_v(d,'cce_research_area')}. We also conducted informal outreach with residents and stakeholders "
        f"in {_v(d,'county')} County who shared their priorities and concerns. This combination of secondary "
        f"research and direct community input shaped our goals and programming approach."
    )

    _add_question(doc, "3b. Who are your community partners, and what role will they play?")
    _add_body(doc,
        f"Our primary community partners include: (1) Cornell Cooperative Extension, which will support "
        f"{_v(d,'cce_role')}; (2) {_v(d,'partner2_name')}, a local or statewide organization focused on "
        f"{_v(d,'partner2_mission')}, which will assist with {_v(d,'partner2_role')}; and "
        f"(3) {_v(d,'partner3_name')}, which will engage through {_v(d,'partner3_engagement')}. Where "
        f"possible, we prioritize partnerships with organizations led by or serving communities "
        f"disproportionately affected by cannabis prohibition. We will formalize agreements through "
        f"{_v(d,'formalization_approach')}."
    )

    # ── SECTION 4 ──────────────────────────────────────────────────────────────
    _add_section_title(doc, "SECTION 4: Goals and Planned Activities")

    _add_question(doc, f"Goal 1: {_v(d,'goal1_title')}")
    _add_body(doc,
        f"Goal 1 is to {_v(d,'goal1_description')}. To achieve this, {_v(d,'business_name')} will "
        f"{_v(d,'goal1_activity_type')}. This activity will take place {_v(d,'goal1_frequency')} and will "
        f"be open to {_v(d,'goal1_target_audience')}. We estimate this activity will reach "
        f"{_v(d,'goal1_participants')} participants per cycle. The event or program will be held at "
        f"{_v(d,'goal1_location')} and will be free of charge or offered at reduced cost to ensure "
        f"accessibility. We will partner with {_v(d,'goal1_partner')} to co-host and promote the activity "
        f"through existing community networks."
    )

    _add_question(doc, f"Goal 2: {_v(d,'goal2_title')}")
    _add_body(doc,
        f"Goal 2 is to {_v(d,'goal2_description')}. {_v(d,'business_name')} will pursue this by "
        f"{_v(d,'goal2_activity')}. Unlike Goal 1, which focuses on {_v(d,'goal1_focus')}, Goal 2 is "
        f"designed to {_v(d,'goal2_purpose')}. To begin, we will conduct {_v(d,'goal2_starting_activity')} "
        f"in partnership with {_v(d,'goal2_partner')}. Based on findings, we will design programming that "
        f"reflects actual community priorities. We anticipate launching this initiative by "
        f"{_v(d,'goal2_launch_date')} with {_v(d,'goal2_participants')} participants in the first cycle."
    )

    if d.get("include_goal3"):
        _add_question(doc, f"Goal 3: {_v(d,'goal3_title')}")
        _add_body(doc,
            f"Goal 3 is to {_v(d,'goal3_description')}. This may include {_v(d,'goal3_activities')}. "
            f"{_v(d,'business_name')} is particularly well-positioned to deliver this because "
            f"{_v(d,'goal3_positioning')}. We will track outcomes through {_v(d,'goal3_tracking')} to "
            f"ensure continuous improvement."
        )

    # ── SECTION 5 ──────────────────────────────────────────────────────────────
    _add_section_title(doc, "SECTION 5: Timeline and Budget")

    _add_question(doc, "5a. Timeline")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for cell, txt in zip(hdr, ["Activity / Goal", "Target Start", "Target Completion / Frequency"]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    timeline_rows = [
        (f"Goal 1: {_v(d,'goal1_title','—')}", _v(d,'tl_g1_start',''), _v(d,'tl_g1_freq','')),
        ("Goal 2 — Phase 1: Needs Assessment / Planning", _v(d,'tl_g2p1_start',''), _v(d,'tl_g2p1_dur','')),
        ("Goal 2 — Phase 2: Program Launch", _v(d,'tl_g2p2_start',''), _v(d,'tl_g2p2_freq','')),
    ]
    if d.get("include_goal3"):
        timeline_rows.append((f"Goal 3: {_v(d,'goal3_title','—')}", _v(d,'tl_g3_start',''), _v(d,'tl_g3_freq','')))

    for act, start, freq in timeline_rows:
        row = tbl.add_row().cells
        row[0].text = act
        row[1].text = start
        row[2].text = freq
        for cell in row:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph()
    _add_question(doc, "5b. Budget")
    budget_period = _v(d, 'budget_period', 'one-year')
    items = []
    for i in (1, 2, 3):
        desc = d.get(f"budget_item{i}", "").strip()
        amt  = d.get(f"budget_amount{i}", "").strip()
        if desc and amt:
            items.append(f"{desc} — ${amt}")
    items_str = "; ".join(items) if items else "[budget items not provided]"
    total_cash = _v(d, "budget_total", "[total not provided]")
    inkind_space = _v(d, "inkind_space", "on-site facilities")
    inkind_space_val = _v(d, "inkind_space_value", "")
    inkind_labor_src = _v(d, "inkind_labor_source", "staff and volunteers")
    inkind_hrs = _v(d, "inkind_labor_hours", "")
    inkind_rate = _v(d, "inkind_labor_rate", "")
    inkind_total = ""
    try:
        inkind_total = f"${float(inkind_hrs.replace(',','')) * float(inkind_rate.replace(',','')):,.0f}"
    except Exception:
        pass
    funding_src = _v(d, "funding_sources", "grant applications and partner contributions")

    labor_str = f"donated labor from {inkind_labor_src}"
    if inkind_hrs and inkind_rate:
        labor_str += f", estimated at {inkind_hrs} hours × ${inkind_rate}/hour"
        if inkind_total:
            labor_str += f" = {inkind_total}"

    _add_body(doc,
        f"The estimated budget for this Community Impact Plan covers a {budget_period} period. "
        f"Cash expenditures include: {items_str}. Estimated total cash budget: approximately "
        f"${total_cash}. In-kind contributions include: use of {inkind_space}"
        + (f" valued at approximately ${inkind_space_val}" if inkind_space_val else "")
        + f"; {labor_str}. We will seek additional support through {funding_src} to offset costs "
        f"as the program grows."
    )

    # ── SECTION 6 ──────────────────────────────────────────────────────────────
    _add_section_title(doc, "SECTION 6: Evaluation and Accountability")

    _add_question(doc, "6a. How will you measure the success of your Community Impact Plan?")
    add_metric = _v(d, 'additional_metric', '')
    qual_method = _v(d, 'qualitative_method', 'participant surveys and feedback forms at events')
    report_freq = _v(d, 'reporting_frequency', 'annually')
    metric_str = (
        "total number of participants reached; number of events or sessions held per year; "
        "number of partner organizations actively engaged"
        + (f"; and {add_metric}" if add_metric not in ("[not provided]", "") else "")
    )
    _add_body(doc,
        f"{_v(d,'business_name')} will track progress toward our Community Impact goals using a "
        f"combination of quantitative and qualitative measures. Quantitative metrics will include: "
        f"{metric_str}. Qualitative data will be collected through {qual_method}. We will compile "
        f"results {report_freq} and share a summary with our community partners and, as required, with OCM."
    )

    _add_question(doc, "6b. How will you adapt your plan based on feedback and results?")
    _add_body(doc,
        f"{_v(d,'business_name')} is committed to responsive, community-centered programming. After each "
        f"activity or program cycle, we will review participant feedback with our core team and partners. "
        f"If {_v(d,'adjustment_condition')}, we will {_v(d,'adjustment_process')}. We will also conduct an "
        f"annual review of the full plan to assess whether goals remain aligned with current community "
        f"priorities. We recognize that community needs evolve, and we are prepared to update this plan "
        f"in consultation with {_v(d,'consultation_partners')}."
    )

    # ── SECTION 7 ──────────────────────────────────────────────────────────────
    if d.get("additional_assets", "").strip() or d.get("long_term_vision", "").strip():
        _add_section_title(doc, "SECTION 7: Additional Information")
        _add_body(doc,
            f"{_v(d,'business_name')} brings the following additional assets and commitments to this "
            f"Community Impact Plan: {_v(d,'additional_assets')}. We believe that cannabis legalization "
            f"presents a generational opportunity to reinvest in communities that bore the heaviest costs "
            f"of prohibition, and we are committed to {_v(d,'long_term_vision')}. We welcome feedback "
            f"from OCM and community stakeholders as we implement this plan."
        )

    # ── Certification ──────────────────────────────────────────────────────────
    _add_section_title(doc, "Certification and Signature")
    _add_body(doc,
        "I certify that the information provided in this Community Impact Plan is accurate and complete "
        "to the best of my knowledge, and that I am authorized to submit this plan on behalf of the "
        "applicant named above."
    )
    doc.add_paragraph()
    for label, key in [
        ("Printed Name:", "contact_name"),
        ("Title:", "title_role"),
        ("Date:", "cert_date"),
    ]:
        _add_field_line(doc, label, _v(d, key, ""))

    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_before = Pt(6)
    sig_label = p_sig.add_run("Signature:  ")
    sig_label.bold = True
    sig_label.font.size = Pt(11)
    p_sig.add_run("_" * 45)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Character counter helper ───────────────────────────────────────────────────
def _char_count(text, limit=4000):
    n = len(text or "")
    pct = n / limit
    if pct >= 1.0:
        return f'<span class="char-over">{n}/{limit} characters — OVER LIMIT</span>'
    if pct >= 0.85:
        return f'<span class="char-warn">{n}/{limit} characters</span>'
    return f'<span style="font-size:0.82rem;color:#555">{n}/{limit} characters</span>'


# ── Page layout ────────────────────────────────────────────────────────────────

st.markdown("""
<div class="disclaimer-box">
<b>This tool generates a draft Community Impact Plan document only.</b>
Review all generated text carefully before submitting to OCM. This is not legal advice.
Always verify current OCM requirements at <a href="https://cannabis.ny.gov" target="_blank">cannabis.ny.gov</a>.
</div>
""", unsafe_allow_html=True)

st.title("📋 NYS OCM Community Impact Plan — Form Builder")
st.caption("Answer the questions in each section to generate a submission-ready CIP document.")

st.info(
    "**How to use this tool:**  \n"
    "Fill in the tabs from left to right — **1 · Applicant Info** through **7 · Additional Info**.  \n"
    "When all sections are complete, click the **✅ Review & Download** tab to generate and download your Word document."
)
st.divider()

if not HAS_DOCX:
    st.error("python-docx is not installed. Run `pip install python-docx` to enable this tool.")
    st.stop()

# ── Session state ──────────────────────────────────────────────────────────────
if "cip_data" not in st.session_state:
    st.session_state.cip_data = {}
d = st.session_state.cip_data

def f(key, default=""):
    return d.get(key, default)

OCM_REGIONS = [
    "Capital Region", "Central New York", "Finger Lakes", "Long Island",
    "Mid-Hudson", "Mohawk Valley", "New York City", "North Country",
    "Southern Tier", "Western New York",
]
LICENSE_TYPES = [
    "Adult-Use Cultivator", "Adult-Use Processor", "Adult-Use Retail Dispensary",
    "Adult-Use Microbusiness", "Adult-Use Delivery", "Adult-Use Cooperative",
    "Registered Organization", "Hemp Grower", "Hemp Processor",
    "Hemp Distributor", "Hemp Retailer", "Nursery",
]

# ── Tabs ───────────────────────────────────────────────────────────────────────
tabs = st.tabs([
    "1 · Applicant Info",
    "2 · Problem & Need",
    "3 · Research & Partners",
    "4 · Goals & Activities",
    "5 · Timeline & Budget",
    "6 · Evaluation",
    "7 · Additional Info",
    "✅ Review & Download",
])

# ── TAB 1: Applicant Information ───────────────────────────────────────────────
with tabs[0]:
    st.markdown("#### Section 1 — Applicant Information")
    st.markdown('<div class="section-note">This section identifies your business and primary contact. All fields are required.</div>', unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        d["business_name"] = st.text_input("Business / Applicant Name *", value=f("business_name"), key="cip_biz_name")
        d["contact_name"] = st.text_input("Primary Contact Name *", value=f("contact_name"), key="cip_contact")
        d["address"] = st.text_input("Street Address *", value=f("address"), key="cip_addr")
        d["city_state_zip"] = st.text_input("City, State, ZIP *", value=f("city_state_zip"), placeholder="e.g. Ithaca, NY 14850", key="cip_csz")
        d["county"] = st.text_input("County *", value=f("county"), placeholder="e.g. Tompkins", key="cip_county")
    with c2:
        lt_idx = LICENSE_TYPES.index(f("license_type")) if f("license_type") in LICENSE_TYPES else 0
        d["license_type"] = st.selectbox("License Type Applied For *", LICENSE_TYPES, index=lt_idx, key="cip_lic")
        d["title_role"] = st.text_input("Title / Role *", value=f("title_role"), placeholder="e.g. Owner, CEO, Operations Manager", key="cip_title")
        ocm_idx = OCM_REGIONS.index(f("ocm_region")) if f("ocm_region") in OCM_REGIONS else 0
        d["ocm_region"] = st.selectbox("OCM Region *", OCM_REGIONS, index=ocm_idx, key="cip_region")
        d["phone"] = st.text_input("Phone *", value=f("phone"), placeholder="e.g. (607) 555-0100", key="cip_phone")
        d["email"] = st.text_input("Email *", value=f("email"), placeholder="name@business.com", key="cip_email")
    st.divider()
    st.caption("➡ When done, click the **2 · Problem & Need** tab above to continue.")

# ── TAB 2: Problem / Need Statement ───────────────────────────────────────────
with tabs[1]:
    st.markdown("#### Section 2 — Problem / Need Statement")
    st.markdown('<div class="section-note">Describe the community need your plan addresses. Focus on your local region — you do not need to address statewide issues.</div>', unsafe_allow_html=True)

    st.markdown("**2a. Your business and the problem it addresses**")
    c1, c2 = st.columns(2)
    with c1:
        d["business_type"] = st.text_input(
            "Type of business",
            value=f("business_type"), placeholder="e.g. farm, dispensary, manufacturer",
            key="cip_biz_type", help="How your business will be described in the document")
        d["affected_backgrounds"] = st.text_input(
            "Backgrounds most affected by prohibition in your area",
            value=f("affected_backgrounds"),
            placeholder="e.g. justice-involved individuals, communities of color, low-income residents",
            key="cip_aff_back")
    with c2:
        d["lack_access_to"] = st.text_input(
            "What do these communities currently lack access to?",
            value=f("lack_access_to"),
            placeholder="e.g. economic opportunities, educational resources, legal support",
            key="cip_lack_access")
        d["region_name"] = st.text_input(
            "Region or area your CIP serves",
            value=f("region_name"),
            placeholder="e.g. Tompkins County, the Southern Tier",
            key="cip_region_name")

    d["unmet_need"] = st.text_area(
        "Specific unmet need your CIP addresses",
        value=f("unmet_need"), height=90,
        placeholder="e.g. lack of agricultural education and workforce development opportunities for justice-involved individuals",
        key="cip_unmet_need")
    st.markdown(_char_count(d["unmet_need"]), unsafe_allow_html=True)

    st.divider()
    st.markdown("**2b. Who is most affected?**")
    d["most_affected_description"] = st.text_area(
        "Describe the individuals or groups most affected",
        value=f("most_affected_description"), height=90,
        placeholder="e.g. justice-involved individuals, residents of historically over-policed neighborhoods, rural residents without economic mobility",
        key="cip_most_aff")
    st.markdown(_char_count(d["most_affected_description"]), unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        d["partners_consulted"] = st.text_input(
            "Community partners or organizations consulted",
            value=f("partners_consulted"),
            placeholder="e.g. local nonprofits, extension offices, community members",
            key="cip_part_consult")
        d["estimated_people"] = st.text_input(
            "Estimated number of people who could benefit",
            value=f("estimated_people"), placeholder="e.g. 500, over 1,000",
            key="cip_est_people")
    with c2:
        d["benefit_area"] = st.text_input(
            "Area where people could benefit",
            value=f("benefit_area"), placeholder="e.g. Tompkins and surrounding counties",
            key="cip_ben_area")
        d["priority_populations"] = st.text_input(
            "Specific populations or zip codes to prioritize",
            value=f("priority_populations"),
            placeholder="e.g. zip codes 14850, 13790; returning citizens",
            key="cip_priority_pop")
    d["outreach_partner"] = st.text_input(
        "Partner who will help reach priority populations",
        value=f("outreach_partner"),
        placeholder="e.g. Tompkins County Reentry Program, local faith communities",
        key="cip_outreach_part")
    st.divider()
    st.caption("➡ When done, click the **3 · Research & Partners** tab above to continue.")

# ── TAB 3: Research & Partnerships ────────────────────────────────────────────
with tabs[2]:
    st.markdown("#### Section 3 — Research, Evidence Base, and Partnerships")
    st.markdown('<div class="section-note">Explain what informed your plan and who your community partners are. OCM values authentic, evidence-based community relationships.</div>', unsafe_allow_html=True)

    st.markdown("**3a. Research and evidence base**")
    d["sources_reviewed"] = st.text_area(
        "Sources reviewed",
        value=f("sources_reviewed"), height=80,
        placeholder="e.g. OCM guidance documents, published literature on cannabis equity, county poverty and arrest rate data, reports from state agencies",
        key="cip_sources")
    st.markdown(_char_count(d["sources_reviewed"]), unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        d["orgs_consulted"] = st.text_input(
            "Organizations or individuals consulted during planning",
            value=f("orgs_consulted"),
            placeholder="e.g. community-based organizations, academic partners, extension services",
            key="cip_orgs_consult")
    with c2:
        d["cce_research_area"] = st.text_input(
            "What did Cornell Cooperative Extension support in your research?",
            value=f("cce_research_area"),
            placeholder="e.g. agricultural education best practices, community engagement strategies",
            key="cip_cce_research")

    st.divider()
    st.markdown("**3b. Community partners**")

    st.markdown("*Partner 1: Cornell Cooperative Extension (required)*")
    d["cce_role"] = st.text_input(
        "CCE's role in your CIP",
        value=f("cce_role"),
        placeholder="e.g. research support, educational programming, needs assessments",
        key="cip_cce_role")

    st.markdown("*Partner 2*")
    c1, c2, c3 = st.columns(3)
    with c1:
        d["partner2_name"] = st.text_input("Organization name", value=f("partner2_name"),
            placeholder="e.g. Ithaca Community Connections", key="cip_p2_name")
    with c2:
        d["partner2_mission"] = st.text_input("Mission / focus area", value=f("partner2_mission"),
            placeholder="e.g. cannabis equity, workforce development", key="cip_p2_mission")
    with c3:
        d["partner2_role"] = st.text_input("Their role in your CIP", value=f("partner2_role"),
            placeholder="e.g. outreach, event co-hosting, referrals", key="cip_p2_role")

    st.markdown("*Partner 3 (Educational Institution)*")
    c1, c2 = st.columns(2)
    with c1:
        d["partner3_name"] = st.text_input("Institution name", value=f("partner3_name"),
            placeholder="e.g. Tompkins Cortland Community College", key="cip_p3_name")
    with c2:
        d["partner3_engagement"] = st.text_input("Type of engagement", value=f("partner3_engagement"),
            placeholder="e.g. student externships, research projects, curriculum partnerships",
            key="cip_p3_engage")

    d["formalization_approach"] = st.text_input(
        "How will you formalize partnerships?",
        value=f("formalization_approach"),
        placeholder="e.g. MOUs, letters of support, co-planning meetings",
        key="cip_formalize")
    st.divider()
    st.caption("➡ When done, click the **4 · Goals & Activities** tab above to continue.")

# ── TAB 4: Goals & Activities ──────────────────────────────────────────────────
with tabs[3]:
    st.markdown("#### Section 4 — Goals and Planned Activities")
    st.markdown('<div class="section-note">List specific, measurable goals and the activities you will carry out. Include estimated participant numbers, frequency, and how activities connect to identified community needs.</div>', unsafe_allow_html=True)

    st.markdown("##### Goal 1")
    d["goal1_title"] = st.text_input("Goal 1 — one-line summary *",
        value=f("goal1_title"),
        placeholder="e.g. Increase access to cannabis agriculture education in Tompkins County",
        key="cip_g1_title")
    d["goal1_description"] = st.text_area(
        "Goal 1 — specific, measurable goal statement",
        value=f("goal1_description"), height=80,
        placeholder="e.g. provide free agricultural education workshops to 50 aspiring cannabis farmers per year",
        key="cip_g1_desc")
    st.markdown(_char_count(d["goal1_description"]), unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1:
        d["goal1_activity_type"] = st.text_input("Activity type",
            value=f("goal1_activity_type"),
            placeholder="e.g. host an annual community workshop series",
            key="cip_g1_act")
        d["goal1_frequency"] = st.text_input("Frequency",
            value=f("goal1_frequency"), placeholder="e.g. quarterly, once per year",
            key="cip_g1_freq")
    with c2:
        d["goal1_target_audience"] = st.text_input("Target audience",
            value=f("goal1_target_audience"),
            placeholder="e.g. justice-involved individuals, low-income residents",
            key="cip_g1_audience")
        d["goal1_participants"] = st.text_input("Estimated participants per cycle",
            value=f("goal1_participants"), placeholder="e.g. 50, 100–150",
            key="cip_g1_part")
    with c3:
        d["goal1_location"] = st.text_input("Location / venue",
            value=f("goal1_location"),
            placeholder="e.g. our farm site, Ithaca Farmers Market pavilion",
            key="cip_g1_loc")
        d["goal1_partner"] = st.text_input("Co-hosting partner",
            value=f("goal1_partner"),
            placeholder="e.g. Cornell Cooperative Extension of Tompkins County",
            key="cip_g1_partner")

    st.divider()
    st.markdown("##### Goal 2")
    d["goal2_title"] = st.text_input("Goal 2 — one-line summary *",
        value=f("goal2_title"),
        placeholder="e.g. Build long-term workforce development capacity for returning citizens",
        key="cip_g2_title")
    d["goal2_description"] = st.text_area(
        "Goal 2 — specific, measurable goal statement",
        value=f("goal2_description"), height=80,
        placeholder="e.g. launch a paid cannabis workforce training cohort for formerly incarcerated individuals",
        key="cip_g2_desc")
    st.markdown(_char_count(d["goal2_description"]), unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        d["goal2_activity"] = st.text_area("Program or initiative description",
            value=f("goal2_activity"), height=80,
            placeholder="e.g. a 6-week hands-on cultivation training program with job placement support",
            key="cip_g2_act")
        d["goal1_focus"] = st.text_input("What does Goal 1 focus on? (brief, for contrast)",
            value=f("goal1_focus"),
            placeholder="e.g. broad community education",
            key="cip_g1_focus")
    with c2:
        d["goal2_purpose"] = st.text_input("What is Goal 2 designed to achieve?",
            value=f("goal2_purpose"),
            placeholder="e.g. build long-term capacity and create ongoing infrastructure",
            key="cip_g2_purpose")
        d["goal2_starting_activity"] = st.text_input("Initial activity to launch Goal 2",
            value=f("goal2_starting_activity"),
            placeholder="e.g. a community needs assessment and listening sessions",
            key="cip_g2_start")
    c1, c2, c3 = st.columns(3)
    with c1:
        d["goal2_partner"] = st.text_input("Partner for Goal 2",
            value=f("goal2_partner"),
            placeholder="e.g. Tompkins County Workforce Development Board",
            key="cip_g2_part")
    with c2:
        d["goal2_launch_date"] = st.text_input("Anticipated launch (quarter and year)",
            value=f("goal2_launch_date"), placeholder="e.g. Q2 2026",
            key="cip_g2_launch")
    with c3:
        d["goal2_participants"] = st.text_input("Participants in first cycle",
            value=f("goal2_participants"), placeholder="e.g. 15–20",
            key="cip_g2_part_n")

    st.divider()
    d["include_goal3"] = st.checkbox("Include Goal 3 (optional)", value=bool(d.get("include_goal3", False)), key="cip_g3_check")
    if d["include_goal3"]:
        st.markdown("##### Goal 3")
        d["goal3_title"] = st.text_input("Goal 3 — one-line summary",
            value=f("goal3_title"),
            placeholder="e.g. Support community health and harm reduction through education",
            key="cip_g3_title")
        d["goal3_description"] = st.text_area("Goal 3 — specific goal statement",
            value=f("goal3_description"), height=80,
            placeholder="e.g. develop a public cannabis consumer safety resource guide",
            key="cip_g3_desc")
        st.markdown(_char_count(d["goal3_description"]), unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1:
            d["goal3_activities"] = st.text_input("Activities included",
                value=f("goal3_activities"),
                placeholder="e.g. hosting student externships, developing a resource guide",
                key="cip_g3_acts")
        with c2:
            d["goal3_positioning"] = st.text_input("Why is your business well-positioned to deliver this?",
                value=f("goal3_positioning"),
                placeholder="e.g. existing partnerships with CCE and community health organizations",
                key="cip_g3_pos")
        with c3:
            d["goal3_tracking"] = st.text_input("How will you track outcomes?",
                value=f("goal3_tracking"),
                placeholder="e.g. sign-in sheets, pre/post surveys, partner feedback",
                key="cip_g3_track")

    st.divider()
    st.caption("➡ When done, click the **5 · Timeline & Budget** tab above to continue.")

# ── TAB 5: Timeline & Budget ───────────────────────────────────────────────────
with tabs[4]:
    st.markdown("#### Section 5 — Timeline and Budget")
    st.markdown('<div class="section-note">Provide a realistic timeline and a general budget breakdown. In-kind contributions (donated time, space, supplies) are acceptable and should be noted separately from cash expenditures.</div>', unsafe_allow_html=True)

    st.markdown("**5a. Timeline**")
    tl1, tl2, tl3 = st.columns(3)
    with tl1:
        st.markdown("*Goal 1*")
        d["tl_g1_start"] = st.text_input("Start", value=f("tl_g1_start"),
            placeholder="e.g. May 2025", key="cip_tl_g1s")
        d["tl_g1_freq"] = st.text_input("Completion / Frequency", value=f("tl_g1_freq"),
            placeholder="e.g. Annual, Ongoing", key="cip_tl_g1f")
    with tl2:
        st.markdown("*Goal 2 — Phase 1: Needs Assessment*")
        d["tl_g2p1_start"] = st.text_input("Start", value=f("tl_g2p1_start"),
            placeholder="e.g. March 2025", key="cip_tl_g2p1s")
        d["tl_g2p1_dur"] = st.text_input("Duration", value=f("tl_g2p1_dur"),
            placeholder="e.g. 8 weeks", key="cip_tl_g2p1d")
    with tl3:
        st.markdown("*Goal 2 — Phase 2: Program Launch*")
        d["tl_g2p2_start"] = st.text_input("Start", value=f("tl_g2p2_start"),
            placeholder="e.g. June 2025", key="cip_tl_g2p2s")
        d["tl_g2p2_freq"] = st.text_input("Frequency / Ongoing cadence", value=f("tl_g2p2_freq"),
            placeholder="e.g. 2 cohorts per year", key="cip_tl_g2p2f")

    if d.get("include_goal3"):
        c1, c2 = st.columns(2)
        with c1:
            d["tl_g3_start"] = st.text_input("Goal 3 — Start", value=f("tl_g3_start"),
                placeholder="e.g. September 2025", key="cip_tl_g3s")
        with c2:
            d["tl_g3_freq"] = st.text_input("Goal 3 — Completion / Frequency", value=f("tl_g3_freq"),
                placeholder="e.g. Annual", key="cip_tl_g3f")

    st.divider()
    st.markdown("**5b. Budget**")
    bp_opts = ["1-year", "2-year", "3-year", "ongoing"]
    bp_idx = bp_opts.index(f("budget_period")) if f("budget_period") in bp_opts else 0
    d["budget_period"] = st.selectbox("Budget period", bp_opts, index=bp_idx, key="cip_bp")

    st.markdown("*Cash expenditures (up to 3 line items)*")
    for i in (1, 2, 3):
        c1, c2 = st.columns([3, 1])
        with c1:
            d[f"budget_item{i}"] = st.text_input(f"Item {i} description",
                value=f(f"budget_item{i}"),
                placeholder="e.g. event supplies and refreshments" if i == 1 else
                            "e.g. printed materials and signage" if i == 2 else
                            "e.g. speaker or facilitator fees",
                key=f"cip_bitem{i}")
        with c2:
            d[f"budget_amount{i}"] = st.text_input(f"Amount ($)",
                value=f(f"budget_amount{i}"), placeholder="e.g. 1,500",
                key=f"cip_bamt{i}")

    d["budget_total"] = st.text_input("Estimated total cash budget ($)",
        value=f("budget_total"), placeholder="e.g. 5,000",
        key="cip_btotal")

    st.markdown("*In-kind contributions*")
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        d["inkind_space"] = st.text_input("Space / facilities",
            value=f("inkind_space"), placeholder="e.g. on-site barn and fields",
            key="cip_ik_space")
    with c2:
        d["inkind_space_value"] = st.text_input("Approx. value ($)",
            value=f("inkind_space_value"), placeholder="e.g. 2,000",
            key="cip_ik_spaceval")
    with c3:
        d["inkind_labor_source"] = st.text_input("Donated labor from",
            value=f("inkind_labor_source"), placeholder="e.g. staff and volunteers",
            key="cip_ik_labsrc")
    with c4:
        d["inkind_labor_hours"] = st.text_input("Hours",
            value=f("inkind_labor_hours"), placeholder="e.g. 200",
            key="cip_ik_hrs")
    with c5:
        d["inkind_labor_rate"] = st.text_input("Rate ($/hr)",
            value=f("inkind_labor_rate"), placeholder="e.g. 20",
            key="cip_ik_rate")

    d["funding_sources"] = st.text_input(
        "Additional funding sources you will pursue",
        value=f("funding_sources"),
        placeholder="e.g. USDA grants, cannabis equity fund applications, partner contributions",
        key="cip_fund_src")
    st.divider()
    st.caption("➡ When done, click the **6 · Evaluation** tab above to continue.")

# ── TAB 6: Evaluation & Accountability ────────────────────────────────────────
with tabs[5]:
    st.markdown("#### Section 6 — Evaluation and Accountability")
    st.markdown('<div class="section-note">OCM expects licensees to track participation, collect community feedback, and adapt programming based on results. Include both quantitative and qualitative methods.</div>', unsafe_allow_html=True)

    st.markdown("**6a. How will you measure success?**")
    st.caption("Standard metrics (total participants, events held, partners engaged) are included automatically.")
    d["additional_metric"] = st.text_input(
        "One additional metric specific to your program",
        value=f("additional_metric"),
        placeholder="e.g. number of individuals placed in employment, educational hours delivered",
        key="cip_add_metric")
    d["qualitative_method"] = st.text_input(
        "Qualitative data collection method",
        value=f("qualitative_method"),
        placeholder="e.g. participant surveys, feedback forms at events, follow-up interviews",
        key="cip_qual_method")
    rf_opts = ["quarterly", "semi-annually", "annually"]
    rf_idx = rf_opts.index(f("reporting_frequency")) if f("reporting_frequency") in rf_opts else 2
    d["reporting_frequency"] = st.selectbox("How often will you compile and report results?",
        rf_opts, index=rf_idx, key="cip_report_freq")

    st.divider()
    st.markdown("**6b. How will you adapt based on feedback?**")
    d["adjustment_condition"] = st.text_input(
        "Condition that would trigger a plan adjustment",
        value=f("adjustment_condition"),
        placeholder="e.g. participation falls below target or feedback indicates an unmet need",
        key="cip_adj_cond")
    d["adjustment_process"] = st.text_area(
        "How you would adjust the plan",
        value=f("adjustment_process"), height=80,
        placeholder="e.g. modify the format, shift the location, adjust timing, or consult with partners to redesign the activity",
        key="cip_adj_proc")
    st.markdown(_char_count(d["adjustment_process"]), unsafe_allow_html=True)
    d["consultation_partners"] = st.text_input(
        "Partners or community representatives you would consult",
        value=f("consultation_partners"),
        placeholder="e.g. our community advisory group and CCE",
        key="cip_consult_part")
    st.divider()
    st.caption("➡ When done, click the **7 · Additional Info** tab above to continue.")

# ── TAB 7: Additional Info + Signature ────────────────────────────────────────
with tabs[6]:
    st.markdown("#### Section 7 — Additional Information (Optional)")
    st.markdown('<div class="section-note">Share anything not captured above: existing community relationships, equity certifications, relevant experience, long-term sustainability vision. This section is optional but strengthens your application.</div>', unsafe_allow_html=True)

    d["additional_assets"] = st.text_area(
        "Additional assets, commitments, or background your business brings",
        value=f("additional_assets"), height=120,
        placeholder="e.g. MWBE certification, 10 years of community service, existing partnerships with local nonprofits, veteran-owned status",
        key="cip_add_assets")
    st.markdown(_char_count(d["additional_assets"]), unsafe_allow_html=True)

    d["long_term_vision"] = st.text_area(
        "Your long-term vision beyond the initial license period",
        value=f("long_term_vision"), height=120,
        placeholder="e.g. growing this program annually, contributing to a regional cannabis equity coalition, mentoring future social equity applicants",
        key="cip_lt_vision")
    st.markdown(_char_count(d["long_term_vision"]), unsafe_allow_html=True)

    st.divider()
    st.markdown("#### Certification")
    st.caption("The signature line will be blank in the downloaded document for wet or digital signature.")
    c1, c2 = st.columns(2)
    with c1:
        st.text_input("Printed Name", value=f("contact_name"), disabled=True,
                      help="Pulled from Section 1", key="cip_cert_name_disp")
        st.text_input("Title", value=f("title_role"), disabled=True,
                      help="Pulled from Section 1", key="cip_cert_title_disp")
    with c2:
        d["cert_date"] = st.text_input("Date of submission",
            value=f("cert_date") or date.today().strftime("%B %d, %Y"),
            key="cip_cert_date")
    st.divider()
    st.caption("➡ When done, click the **✅ Review & Download** tab above to generate your document.")

# ── TAB 8: Review & Download ───────────────────────────────────────────────────
with tabs[7]:
    st.markdown("#### Review & Download")

    # Completeness check
    required = {
        "Business Name": d.get("business_name", ""),
        "License Type": d.get("license_type", ""),
        "Contact Name": d.get("contact_name", ""),
        "County": d.get("county", ""),
        "OCM Region": d.get("ocm_region", ""),
        "Business Type": d.get("business_type", ""),
        "Unmet Need": d.get("unmet_need", ""),
        "Goal 1 Title": d.get("goal1_title", ""),
        "Goal 1 Description": d.get("goal1_description", ""),
        "Goal 2 Title": d.get("goal2_title", ""),
        "Goal 2 Description": d.get("goal2_description", ""),
        "CCE Role": d.get("cce_role", ""),
    }
    missing = [k for k, v in required.items() if not v.strip()]

    if missing:
        st.warning(f"The following required fields are not yet filled in: **{', '.join(missing)}**. "
                   "You can still download a draft, but blank fields will appear as '[not provided]'.")
    else:
        st.success("✅ All required fields are complete. Your document is ready to download.")

    st.divider()

    # Preview key info
    st.markdown("**Document Preview — Key Information**")
    prev_cols = st.columns(3)
    with prev_cols[0]:
        st.markdown(f"**Business:** {d.get('business_name','—')}")
        st.markdown(f"**License Type:** {d.get('license_type','—')}")
        st.markdown(f"**County / Region:** {d.get('county','—')} / {d.get('ocm_region','—')}")
    with prev_cols[1]:
        st.markdown(f"**Goal 1:** {d.get('goal1_title','—')}")
        st.markdown(f"**Goal 2:** {d.get('goal2_title','—')}")
        if d.get("include_goal3"):
            st.markdown(f"**Goal 3:** {d.get('goal3_title','—')}")
    with prev_cols[2]:
        st.markdown(f"**Budget Period:** {d.get('budget_period','—')}")
        st.markdown(f"**Total Cash Budget:** ${d.get('budget_total','—')}")
        st.markdown(f"**Submission Date:** {d.get('cert_date','—')}")

    st.divider()
    st.markdown("### 📄 Generate Your Community Impact Plan Document")
    st.markdown(
        "Click the button below to build your Word document from all the information you entered. "
        "The file will download immediately — open it in Microsoft Word or Google Docs to review and sign."
    )

    docx_bytes = build_cip_docx(d)
    filename = f"CIP_{d.get('business_name','draft').replace(' ','_')}_{d.get('cert_date','').replace(' ','_').replace(',','')}.docx"

    st.download_button(
        "📄 Generate & Download Community Impact Plan (.docx)",
        data=docx_bytes,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        type="primary",
    )

    st.caption(
        "The downloaded file is a Microsoft Word (.docx) document. Review all text carefully before "
        "submitting to OCM. Open-ended narrative responses must be 4,000 characters or fewer per OCM instructions."
    )
